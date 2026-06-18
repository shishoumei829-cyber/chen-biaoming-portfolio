from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from pathlib import Path


OUT = Path(r"C:\Users\SHIKIMORI\Desktop\简历_新版.docx")
BACKUP = Path(r"C:\Users\SHIKIMORI\Desktop\简历_新版_原版备份.docx")


NAME_FONT = "Microsoft YaHei"
BODY_FONT = "Microsoft YaHei"
ACCENT = RGBColor(38, 70, 83)
MUTED = RGBColor(90, 96, 104)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, bottom=90, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9DEE3", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_run(run, size=9, bold=False, color=None, font=BODY_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def p_style(p, before=0, after=2, line=1.03, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align


def add_text(p, text, size=9, bold=False, color=None, font=BODY_FONT):
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color, font=font)
    return r


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p_style(p, before=5, after=2)
    add_text(p, text, size=10.5, bold=True, color=ACCENT)
    pbdr = p._p.get_or_add_pPr().find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        p._p.get_or_add_pPr().append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "9FB3BF")
    pbdr.append(bottom)


def add_bullet(doc, label, text):
    p = doc.add_paragraph(style="List Bullet")
    p_style(p, before=0, after=1.5, line=1.02)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    add_text(p, label, size=8.7, bold=True)
    add_text(p, " " + text, size=8.7)


def add_project(doc, title, meta, bullets):
    p = doc.add_paragraph()
    p_style(p, before=2, after=1)
    add_text(p, title, size=9.7, bold=True, color=RGBColor(20, 35, 43))
    add_text(p, "    " + meta, size=8.2, color=MUTED)
    for label, text in bullets:
        add_bullet(doc, label, text)


def add_hyperlink(paragraph, text, url, color="2E74B5"):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build():
    if OUT.exists() and not BACKUP.exists():
        BACKUP.write_bytes(OUT.read_bytes())

    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.03

    bullet = styles["List Bullet"]
    bullet.font.name = BODY_FONT
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    bullet.font.size = Pt(8.7)
    bullet.paragraph_format.space_after = Pt(1.5)

    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    header.columns[0].width = Inches(3.0)
    header.columns[1].width = Inches(4.1)
    set_table_borders(header, color="FFFFFF", size="0")
    for row in header.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=0, bottom=20, start=0, end=0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left, right = header.rows[0].cells
    p = left.paragraphs[0]
    p_style(p, after=0)
    add_text(p, "陈彪明", size=22, bold=True, color=RGBColor(17, 24, 39), font=NAME_FONT)
    p = left.add_paragraph()
    p_style(p, after=0)
    add_text(p, "用户体验设计师 · 产品设计师 · 产品经理", size=9.4, color=ACCENT)

    p = right.paragraphs[0]
    p_style(p, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(p, "桂林 · 广西", size=8.7, color=MUTED)
    p = right.add_paragraph()
    p_style(p, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(p, "电话：18078193563    邮箱：1310258735@qq.com", size=8.7, color=MUTED)
    p = right.add_paragraph()
    p_style(p, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(p, "作品集：", size=8.7, color=MUTED)
    add_hyperlink(p, "chen-biaoming-portfolio", "https://shishoumei829-cyber.github.io/chen-biaoming-portfolio/")

    p = doc.add_paragraph()
    p_style(p, before=2, after=5)
    add_text(p, "工业设计背景，聚焦 UX / 产品方向。擅长从用户痛点出发完成需求定义、交互设计与高保真产品交付。", size=9.2)

    add_section_heading(doc, "教育背景")
    p = doc.add_paragraph()
    p_style(p, after=1)
    add_text(p, "桂林电子科技大学 · 工业设计", size=9.3, bold=True)
    add_text(p, "    在读（大二）", size=8.7, color=MUTED)
    p = doc.add_paragraph()
    p_style(p, after=3)
    add_text(p, "暂无实习经历，持续寻找 UX / 产品方向实习机会", size=8.8, color=MUTED)

    add_section_heading(doc, "项目经历")
    add_project(doc, "MIRAGE · 现实转译器", "2026.03 — 05 · 独立负责", [
        ("项目背景", "针对日常通勤场景中对街景美感钝化、屏幕分散注意力的问题，独立负责一款以「现实转译」为核心的 AR 眼镜产品概念设计。"),
        ("设计决策", "运用空间计算定义「现实转译」交互范式，替代传统滤镜叠加；设计空间对齐助手 MIA，解决 2D 弹窗式 AI 破坏沉浸感与信任感的痛点；建立感官冗余机制，平衡视觉美化与环境安全识别。"),
        ("交付成果", "输出完整产品定义文档与用户场景矩阵，交付 MIA 核心交互逻辑说明，并完成硬件人机工学与材料规格方案。"),
    ])
    add_project(doc, "AMADEUS · 数字生命实验系统", "2026.03 — 05 · 独立负责", [
        ("项目背景", "针对独居群体情感陪伴需求，独立负责一款基于大语言模型的 AI 陪伴系统设计。"),
        ("设计决策", "运用 PAD 情绪矩阵构建动态行为仲裁机制；设计「记忆宫殿」语义检索交互，解决 AI 对话长期记忆断层痛点。"),
        ("交付成果", "绘制 4 个核心场景的用户旅程图（User Journey），交付高保真全链路交互设计稿，并完成基于开发者透明模式的界面可用性测试（Usability Testing）。"),
    ])
    add_project(doc, "墨舟 · 时空文化导游", "2026.03 — 05 · 独立负责", [
        ("需求定义", "洞察轻度文化游用户的「即时探索」痛点，提出「到达即阅读」的产品概念。"),
        ("交互优化", "采用基于地理位置（LBS）的瞬间触发机制；设计轻量化「轻量文化卡」与「近处回响弹层」，减少用户操作路径，避免传统百科式界面的信息过载。"),
        ("输出交付", "独立完成探索、发现、雅藏 3 大核心流程的交互原型设计，与雅集排行等高保真界面交付。"),
    ])

    add_section_heading(doc, "技能")
    skills = doc.add_table(rows=4, cols=2)
    skills.alignment = WD_TABLE_ALIGNMENT.CENTER
    skills.autofit = False
    skills.columns[0].width = Inches(1.25)
    skills.columns[1].width = Inches(5.9)
    set_table_borders(skills, color="E3E7EA", size="4")
    rows = [
        ("专业技能", "用户研究、交互设计、原型绘制、产品叙事、需求分析、全链路设计"),
        ("设计与界面", "Stitch · Rhinoceros · SolidWorks"),
        ("视觉与影像", "Photoshop · DaVinci Resolve"),
        ("AI 效能工具", "Cursor · Copilot · ChatGPT（具备利用 AI 辅助快速构建原型与前端复现的能力）"),
    ]
    for row, (k, v) in zip(skills.rows, rows):
        for cell in row.cells:
            set_cell_margins(cell, top=80, bottom=80, start=110, end=110)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], "F3F6F7")
        p = row.cells[0].paragraphs[0]
        p_style(p, after=0)
        add_text(p, k, size=8.5, bold=True, color=ACCENT)
        p = row.cells[1].paragraphs[0]
        p_style(p, after=0)
        add_text(p, v, size=8.5)

    footer = section.footer.paragraphs[0]
    p_style(footer, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(footer, "陈彪明 · 简历 · 2026", size=7.5, color=MUTED)

    doc.save(OUT)


if __name__ == "__main__":
    build()
